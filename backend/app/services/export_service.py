from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from pathlib import Path
from datetime import datetime


EXPORT_DIR = "exports"

Path(EXPORT_DIR).mkdir(exist_ok=True)

# Register professional TTF fonts if available
FONT_REGISTERED = False

try:
    # Try to register a professional serif font
    pdfmetrics.registerFont(TTFont("Georgia", "C:/Windows/Fonts/Georgia.ttf"))
    pdfmetrics.registerFont(TTFont("Georgia-Bold", "C:/Windows/Fonts/georgiab.ttf"))
    pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/Arial.ttf"))
    pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
    pdfmetrics.registerFont(TTFont("Arial-Italic", "C:/Windows/Fonts/ariali.ttf"))
    pdfmetrics.registerFont(TTFont("Arial-BoldItalic", "C:/Windows/Fonts/arialbi.ttf"))
    FONT_REGISTERED = True
except Exception:
    # Fall back to built-in fonts
    FONT_REGISTERED = False


# -----------------------------------
# Helper Formatter
# -----------------------------------
def format_items(items):
    """Clean and format extracted items for display."""
    formatted = []

    for item in items:
        # If dict with confidence
        if isinstance(item, dict):
            value = item.get("value", "")

            confidence = item.get(
                "confidence",
                None
            )

            if confidence is not None:
                formatted.append(
                    f"{value} "
                    f"(Confidence: {confidence:.0%})"
                )

            else:
                formatted.append(value)

        else:
            # Clean the string - remove markdown artifacts
            item_str = str(item).strip()

            # Remove markdown code block markers
            item_str = item_str.replace("```json", "").replace("```", "")
            item_str = item_str.replace("**", "").replace("*", "")

            formatted.append(item_str)

    return formatted


def clean_text(text: str) -> str:
    """Remove markdown and formatting artifacts."""
    text = text.replace("```json", "").replace("```", "")
    text = text.replace("**", "").replace("*", "")

    text = text.replace("`", "")
    return text.strip()


def parse_markdown_to_paragraphs(text: str, styles) -> list:
    """Parse markdown text into a list of (style_name, content) tuples for PDF."""
    elements = []
    lines = text.split("\n")
    in_list = False

    body_style = styles.get("CustomBody", styles["BodyText"])
    heading_style = styles.get("CustomHeading", styles["Heading2"])
    bullet_style = styles.get("CustomBullet", styles["BodyText"])

    # Custom styles for parsed markdown
    h1_style = ParagraphStyle(
        "MarkdownH1",
        parent=heading_style,
        fontSize=16,
        textColor=colors.HexColor("#1a1a2e"),
        spaceBefore=12,
        spaceAfter=6,
    )

    h2_style = ParagraphStyle(
        "MarkdownH2",
        parent=heading_style,
        fontSize=14,
        textColor=colors.HexColor("#1a1a2e"),
        spaceBefore=10,
        spaceAfter=6,
    )

    h3_style = ParagraphStyle(
        "MarkdownH3",
        parent=heading_style,
        fontSize=12,
        textColor=colors.HexColor("#333333"),
        spaceBefore=8,
        spaceAfter=4,
    )

    plain_style = ParagraphStyle(
        "MarkdownPlain",
        parent=body_style,
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
    )

    list_bullet_style = ParagraphStyle(
        "MarkdownListBullet",
        parent=bullet_style,
        fontSize=10,
        leading=14,
        leftIndent=20,
        spaceAfter=3,
        bulletIndent=10,
    )

    def split_heading_content(line):
        """Split heading from content if on same line like '### Heading Content'."""
        # Find where the heading ends and content begins
        # Pattern: after heading marker, find first period+space that's NOT at end
        # followed by capital letter or common content words
        for i, c in enumerate(line):
            if c == '.' and i > 30 and i < len(line) - 20:
                # Possible split point - check what follows
                rest = line[i+1:].strip()
                if rest and rest[0].isupper():
                    return line[:i], line[i+1:].strip()
        return None, line  # No split found

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            in_list = False
            continue

        # H1 heading (# Header)
        if stripped.startswith("# "):
            heading, rest = split_heading_content(stripped[2:])
            if heading:
                elements.append((None, "skip"))
                elements.append(("h1", heading))
                if rest:
                    elements.append(("body", rest))
            else:
                elements.append((None, "skip"))
                elements.append(("h1", stripped[2:].strip()))
        # H2 heading (## Header)
        elif stripped.startswith("## "):
            heading, rest = split_heading_content(stripped[3:])
            if heading:
                elements.append((None, "skip"))
                elements.append(("h2", heading))
                if rest:
                    elements.append(("body", rest))
            else:
                elements.append((None, "skip"))
                elements.append(("h2", stripped[3:].strip()))
        # H3 heading (### Header)
        elif stripped.startswith("### "):
            heading, rest = split_heading_content(stripped[4:])
            if heading:
                elements.append((None, "skip"))
                elements.append(("h3", heading))
                if rest:
                    elements.append(("body", rest))
            else:
                elements.append((None, "skip"))
                elements.append(("h3", stripped[4:].strip()))
        # Bullet list item (* or • or -)
        elif stripped.startswith(("* ", "• ", "- ")):
            content = stripped[2:].strip()
            if content:
                elements.append(("list_bullet", content))
                in_list = True
        # Numbered list (1. 2. etc)
        elif stripped and stripped[0].isdigit() and ". " in stripped[:4]:
            content = stripped[stripped.index(". ") + 2:].strip()
            if content:
                elements.append(("list_bullet", content))
                in_list = True
        # Bold text line (standalone **text**)
        elif stripped.startswith("**") and stripped.endswith("**"):
            elements.append(("bold", stripped[2:-2].strip()))
        # Regular paragraph
        else:
            # Clean markdown artifacts but preserve text
            clean = stripped
            # Remove inline markdown
            clean = clean.replace("**", "").replace("*", "")
            clean = clean.replace("`", "").replace("```", "")
            if clean:
                elements.append(("body", clean))

    return elements


def render_markdown_to_elements(text: str, styles, body_font="Times") -> list:
    """Convert markdown text to PDF elements with proper styling."""
    elements = []
    parsed = parse_markdown_to_paragraphs(text, styles)

    # Define styles
    h2_style = ParagraphStyle(
        "ParsedH2",
        fontName="Helvetica-Bold",
        fontSize=14,
        spaceBefore=14,
        spaceAfter=8,
        textColor=colors.HexColor("#1a1a2e"),
        leading=18,
    )

    h3_style = ParagraphStyle(
        "ParsedH3",
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceBefore=10,
        spaceAfter=6,
        textColor=colors.HexColor("#333333"),
        leading=16,
    )

    body_style = ParagraphStyle(
        "ParsedBody",
        fontName=body_font,
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )

    bullet_style = ParagraphStyle(
        "ParsedBullet",
        fontName=body_font,
        fontSize=10,
        leading=14,
        leftIndent=25,
        spaceAfter=4,
        bulletIndent=10,
    )

    bold_style = ParagraphStyle(
        "ParsedBold",
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )

    for style_name, content in parsed:
        if style_name == "skip":
            elements.append(Spacer(1, 6))
        elif style_name == "h2":
            elements.append(Paragraph(content, h2_style))
        elif style_name == "h3":
            elements.append(Paragraph(content, h3_style))
        elif style_name == "body":
            elements.append(Paragraph(content, body_style))
        elif style_name == "list_bullet":
            elements.append(Paragraph(f"\u2022 {content}", bullet_style))
        elif style_name == "bold":
            elements.append(Paragraph(content, bold_style))

    return elements


def set_run_font(run, name, size, bold=False, italic=False):
    """Set font on a docx run with cross-platform font name."""
    run.font.name = name
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    # Also set the East Asian font for wide character support
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)


# -----------------------------------
# DOCX EXPORT
# -----------------------------------
def generate_docx_report(report_data):
    document = Document()

    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Set default paragraph font to Calibri 11pt
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title - Cambria 24pt Bold
    title = document.add_heading("RFP Analysis Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Cambria"
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Subtitle - Calibri 12pt
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Generated Executive Analysis")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Add date - Calibri 10pt
    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    document.add_paragraph()

    # SECTION DEFINITIONS
    section_defs = [
        ("Executive Summary", "summary"),
        ("Scope of Work", "scope_of_work"),
        ("Deliverables", "deliverables"),
        ("Objectives", "objectives"),
        ("Deadlines", "deadlines"),
        ("Staffing Requirements", "staffing_requirements"),
        ("Compliance Items", "compliance_items"),
    ]

    def add_section(title_text, items, is_first=False):
        if not is_first:
            document.add_paragraph()

        # Section heading - Cambria 14pt Bold
        heading = document.add_heading(title_text, level=2)
        for run in heading.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        formatted_items = format_items(items)

        if not formatted_items:
            para = document.add_paragraph("No information extracted.")
            para.style = document.styles["Normal"]
            run = para.runs[0] if para.runs else para.add_run("")
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        else:
            for item in formatted_items:
                para = document.add_paragraph(item, style="List Bullet")
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

    # Executive Summary - special handling
    summary_items = report_data.get("summary", [])

    if summary_items:
        document.add_heading("Executive Summary", level=2)
        for run in document.paragraphs[-1].runs:
            run.font.name = "Cambria"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Handle executive summary as one block of text
        first = True
        for item in summary_items:
            if isinstance(item, dict):
                text = item.get("value", "")
            else:
                text = str(item)

            text = clean_text(text)

            if text:
                para = document.add_paragraph()
                run = para.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                if not first:
                    para.paragraph_format.space_before = Pt(0)
                first = False
    else:
        document.add_heading("Executive Summary", level=2)
        for run in document.paragraphs[-1].runs:
            run.font.name = "Cambria"
            run.font.size = Pt(14)
            run.font.bold = True
        para = document.add_paragraph("No information extracted.")
        run = para.runs[0] if para.runs else para.add_run("")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Other sections
    first = True
    for title_text, key in section_defs[1:]:
        items = report_data.get(key, [])
        add_section(title_text, items, first)
        if first:
            first = False

    # Footer note
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("This report was automatically generated by RFP Analysis Agent.")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    file_path = f"{EXPORT_DIR}/rfp_report.docx"

    document.save(file_path)

    return file_path


# -----------------------------------
# PDF EXPORT
# -----------------------------------
def generate_pdf_report(report_data):
    file_path = f"{EXPORT_DIR}/rfp_report.pdf"

    doc = SimpleDocTemplate(
        file_path,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()

    # Choose fonts based on availability
    if FONT_REGISTERED:
        title_font = "Georgia"
        heading_font = "Arial"
        body_font = "Georgia"
        bullet_font = "Georgia"
    else:
        # Built-in professional fonts
        title_font = "Times-Bold"
        heading_font = "Helvetica-Bold"
        body_font = "Times"
        bullet_font = "Times"

    # Custom styles with professional fonts
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName=title_font,
        fontSize=22,
        spaceAfter=6,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1a1a2e"),
        leading=26
    )

    subtitle_style = ParagraphStyle(
        "CustomSubtitle",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=12,
        spaceAfter=2,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555")
    )

    date_style = ParagraphStyle(
        "CustomDate",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=10,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#888888")
    )

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontName=heading_font,
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#1a1a2e"),
        leading=18
    )

    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=10,
        leading=15,
        alignment=TA_JUSTIFY
    )

    bullet_style = ParagraphStyle(
        "CustomBullet",
        parent=styles["BodyText"],
        fontName=bullet_font,
        fontSize=10,
        leading=14,
        leftIndent=20,
        spaceAfter=4
    )

    empty_style = ParagraphStyle(
        "EmptyStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#999999")
    )

    elements = []

    # Title
    elements.append(Paragraph("RFP Analysis Report", title_style))
    elements.append(Paragraph("AI-Generated Executive Analysis", subtitle_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", date_style))

    # Horizontal rule
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#ddd"), spaceAfter=12))

    # SECTION DEFINITIONS
    section_defs = [
        ("Executive Summary", "summary"),
        ("Scope of Work", "scope_of_work"),
        ("Deliverables", "deliverables"),
        ("Objectives", "objectives"),
        ("Deadlines", "deadlines"),
        ("Staffing Requirements", "staffing_requirements"),
        ("Compliance Items", "compliance_items"),
    ]

    first = True
    for section_title, key in section_defs:
        if not first:
            elements.append(Spacer(1, 12))

        elements.append(Paragraph(section_title, heading_style))

        items = report_data.get(key, [])
        formatted_items = format_items(items)

        if not formatted_items:
            elements.append(
                Paragraph("No information extracted.", empty_style)
            )
        else:
            # Check if this is markdown content (Executive Summary)
            if key == "summary" and len(formatted_items) == 1:
                # Try to render as markdown
                md_text = formatted_items[0]
                if "## " in md_text or "### " in md_text or "* " in md_text or "• " in md_text:
                    # Contains markdown - render properly
                    md_elements = render_markdown_to_elements(
                        md_text, styles, body_font=body_font
                    )
                    elements.extend(md_elements)
                else:
                    # Plain text - justify it
                    para_style = ParagraphStyle(
                        "SummaryBody",
                        parent=body_style,
                        alignment=TA_JUSTIFY,
                    )
                    elements.append(Paragraph(md_text, para_style))
            else:
                for item in formatted_items:
                    elements.append(
                        Paragraph(f"\u2022 {item}", bullet_style)
                    )

        first = False

    # Footer
    elements.append(Spacer(1, 24))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ddd")))
    elements.append(Spacer(1, 6))
    footer_style = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#AAAAAA")
    )
    elements.append(
        Paragraph("This report was automatically generated by RFP Analysis Agent.", footer_style)
    )

    doc.build(elements)

    return file_path
